import re
from dataclasses import dataclass
from io import BytesIO
from typing import Iterable
import base64

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from fastapi.concurrency import run_in_threadpool

from src.html_to_docx.exceptions import HTMLConversionError


@dataclass(frozen=True)
class InlineStyle:
    bold: bool = False
    italic: bool = False
    underline: bool = False


BLOCK_TAGS = {
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "p",
    "div",
    "section",
    "article",
    "header",
    "footer",
    "table",
    "ul",
    "ol",
    "li",
    "img",
    "hr",
}


# Regras CSS por classe (parse simplificado de <style>).
_CSS_CLASS_STYLES: dict[str, dict[str, str]] = {}


def _set_default_document_style(doc: Document, font_name: str = "Calibri", font_size_pt: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    # aproxima o CSS do email: line-height: 1.5
    style.paragraph_format.line_spacing = 1.5


def _normalize_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _parse_css_class_rules(css_text: str) -> dict[str, dict[str, str]]:
    """
    Parser simples de CSS: captura regras do tipo `.classe { a:b; c:d }`.
    Suficiente pra HTML de e-mail (logo-img, header-title, etc.).
    """
    out: dict[str, dict[str, str]] = {}
    if not css_text:
        return out

    # remove comentários
    css_text = re.sub(r"/\*[\s\S]*?\*/", "", css_text)
    # captura blocos `.foo { ... }` (sem suportar seletores complexos)
    for m in re.finditer(r"\.(?P<cls>[A-Za-z0-9_-]+)\s*\{(?P<body>[^}]+)\}", css_text):
        cls = m.group("cls")
        body = m.group("body")
        out[cls] = _parse_style_attr(body)
    return out


def _style_map_for(tag: Tag) -> dict[str, str]:
    """
    Merge de estilos: primeiro CSS por class, depois style inline (inline ganha).
    """
    merged: dict[str, str] = {}
    for cls in (tag.get("class") or []):
        merged.update(_CSS_CLASS_STYLES.get(str(cls), {}))
    merged.update(_parse_style_attr(tag.get("style") or ""))
    return merged


def _parse_style_attr(style: str) -> dict[str, str]:
    """
    Converte "a:b; c:d" em {"a": "b", "c": "d"} (lowercase).
    """
    out: dict[str, str] = {}
    for part in (style or "").split(";"):
        if ":" not in part:
            continue
        k, v = part.split(":", 1)
        k = k.strip().lower()
        v = v.strip()
        if k:
            out[k] = v
    return out


def _extract_px(style_map: dict[str, str], prop: str) -> int | None:
    val = style_map.get(prop.lower())
    if not val:
        return None
    m = re.search(r"(\d+(?:\.\d+)?)\s*px", val, flags=re.IGNORECASE)
    if not m:
        return None
    try:
        return int(float(m.group(1)))
    except Exception:
        return None


def _closest_text_align_right(tag: Tag) -> bool:
    """
    Sobe a árvore procurando style="text-align: right" (comum em HTML de e-mail).
    """
    cur: Tag | None = tag
    while isinstance(cur, Tag):
        style_map = _style_map_for(cur)
        align = (style_map.get("text-align") or "").lower()
        if align == "right":
            return True
        cur = cur.parent if isinstance(cur.parent, Tag) else None
    return False


def _remove_table_borders(table) -> None:
    tbl = table._tbl  # noqa: SLF001
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _is_real_table(table: Tag) -> bool:
    """
    Heurística: muitos HTMLs de e-mail usam <table> apenas pra layout.
    Só cria tabela no DOCX quando parece realmente tabular.
    """
    if not isinstance(table, Tag) or table.name != "table":
        return False

    classes = set(table.get("class") or [])
    if "data-table" in classes:
        return True

    if table.find("th") is not None:
        return True

    border = table.get("border")
    if border and str(border).strip() not in ("0", ""):
        return True

    return False


def _iter_child_nodes(tag: Tag) -> Iterable[object]:
    for child in tag.children:
        if isinstance(child, NavigableString):
            if str(child).strip():
                yield child
        elif isinstance(child, Tag):
            yield child


def _add_hyperlink(paragraph, url: str, text: str, style: InlineStyle) -> None:
    """
    python-docx não tem API de hyperlink de alto nível; precisamos inserir XML.
    """
    if not url:
        run = paragraph.add_run(text)
        run.bold = style.bold
        run.italic = style.italic
        run.underline = style.underline
        return

    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    if style.bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
    if style.italic:
        i = OxmlElement("w:i")
        r_pr.append(i)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    r_pr.append(color)

    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)  # noqa: SLF001 (python-docx API interna)


def _add_image(container, img: Tag) -> None:
    src = (img.get("src") or "").strip()
    if not src:
        return

    # Caso comum em emails: data:image/png;base64,...
    m = re.match(r"^data:image/(?P<fmt>[^;]+);base64,(?P<b64>.+)$", src, flags=re.IGNORECASE | re.DOTALL)
    if m:
        b64 = m.group("b64")
        try:
            raw = base64.b64decode(b64, validate=False)
        except Exception:
            return

        bio = BytesIO(raw)
        # tenta respeitar max-width do HTML (px -> inches usando 96dpi como aproximação)
        style_map = _style_map_for(img)
        max_w_px = _extract_px(style_map, "max-width")
        if not max_w_px:
            max_w_px = _extract_px(style_map, "width")
        if not max_w_px:
            # atributo HTML width="150"
            try:
                max_w_px = int(str(img.get("width") or "").strip() or "0") or None
            except Exception:
                max_w_px = None
        width = Inches(max_w_px / 96.0) if max_w_px else None
        # insere em parágrafo (pra conseguir alinhar)
        p = container.add_paragraph()
        if _closest_text_align_right(img):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        try:
            if width:
                run.add_picture(bio, width=width)
            else:
                run.add_picture(bio)
        except Exception:
            return
        return

    # URL/arquivo: python-docx precisa de path/stream; pra URL seria necessário baixar.
    # Mantemos simples e ignoramos (mas poderia ser implementado depois).
    return


def _process_inline(paragraph, node: object, style: InlineStyle) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = style.bold
        run.italic = style.italic
        run.underline = style.underline
        return

    if not isinstance(node, Tag):
        return

    name = node.name.lower()

    if name in ("strong", "b"):
        new_style = InlineStyle(bold=True or style.bold, italic=style.italic, underline=style.underline)
        for child in _iter_child_nodes(node):
            _process_inline(paragraph, child, new_style)
        return

    if name in ("em", "i"):
        new_style = InlineStyle(bold=style.bold, italic=True or style.italic, underline=style.underline)
        for child in _iter_child_nodes(node):
            _process_inline(paragraph, child, new_style)
        return

    if name == "u":
        new_style = InlineStyle(bold=style.bold, italic=style.italic, underline=True or style.underline)
        for child in _iter_child_nodes(node):
            _process_inline(paragraph, child, new_style)
        return

    if name == "br":
        paragraph.add_run().add_break()
        return

    if name == "a":
        url = (node.get("href") or "").strip()
        text = _normalize_ws(node.get_text(" ", strip=True)) or url
        _add_hyperlink(paragraph, url, text, style)
        return

    if name == "img":
        # Imagem dentro de inline: ignora (tratamos <img> como bloco quando aparece como nó do container)
        return

    # default: span, etc.
    for child in _iter_child_nodes(node):
        _process_inline(paragraph, child, style)


def _process_list(container, lst: Tag, ordered: bool) -> None:
    style_name = "List Number" if ordered else "List Bullet"
    for li in lst.find_all("li", recursive=False):
        p = container.add_paragraph(style=style_name)
        for child in _iter_child_nodes(li):
            if isinstance(child, Tag) and child.name.lower() in BLOCK_TAGS and child.name.lower() not in ("strong", "b", "em", "i", "u", "a", "br", "span"):
                # Se houver bloco dentro do <li>, processa como bloco (após o prefixo do item).
                _process_block(container, child)
            else:
                _process_inline(p, child, InlineStyle())


def _process_table(container, table_tag: Tag) -> None:
    # evita "duplicar" conteúdo quando existem tabelas aninhadas (muito comum em e-mails)
    rows = [tr for tr in table_tag.find_all("tr") if tr.find_parent("table") is table_tag]
    if not rows:
        return

    # calcula número máximo de colunas (respeitando HTML "desbalanceado")
    max_cols = 0
    row_cells = []
    for r in rows:
        cells = r.find_all(["td", "th"], recursive=False)
        row_cells.append(cells)
        max_cols = max(max_cols, len(cells))

    if max_cols == 0:
        return

    tbl = container.add_table(rows=len(rows), cols=max_cols)
    tbl.style = "Table Grid"

    for i, cells in enumerate(row_cells):
        for j in range(max_cols):
            cell = tbl.cell(i, j)
            if j >= len(cells):
                cell.text = ""
                continue
            txt = cells[j].get_text(" ", strip=True)
            cell.text = _normalize_ws(txt)


def _process_layout_table_as_docx_table(container, table_tag: Tag) -> None:
    """
    Tabelas de layout em e-mails carregam a estrutura (colunas/alinhamentos).
    Aqui criamos uma tabela do Word sem bordas e processamos conteúdo por célula.
    """
    rows = [tr for tr in table_tag.find_all("tr") if tr.find_parent("table") is table_tag]
    if not rows:
        return

    row_cells: list[list[Tag]] = []
    max_cols = 0
    for r in rows:
        cells = r.find_all(["td", "th"], recursive=False)
        row_cells.append(cells)
        max_cols = max(max_cols, len(cells))

    if max_cols == 0:
        return

    tbl = container.add_table(rows=len(rows), cols=max_cols)
    _remove_table_borders(tbl)
    tbl.autofit = True

    for i, cells in enumerate(row_cells):
        for j in range(max_cols):
            docx_cell = tbl.cell(i, j)
            # limpa parágrafo padrão
            if docx_cell.paragraphs:
                docx_cell.paragraphs[0].text = ""
            if j >= len(cells):
                continue
            _process_container(docx_cell, cells[j])


def _process_heading(container, tag: Tag) -> None:
    level = int(tag.name[1]) if tag.name and tag.name.startswith("h") and tag.name[1:].isdigit() else 2
    level = max(1, min(6, level))
    text = _normalize_ws(tag.get_text(" ", strip=True))
    if text:
        # python-docx expõe Document() como função fábrica; checamos por método.
        p = container.add_heading(text, level=level) if hasattr(container, "add_heading") else container.add_paragraph(text)
        # aproxima estilo de email: h1/h2 geralmente maior e pode ter cor azul
        if p.runs:
            r = p.runs[0]
            if level == 1:
                r.font.size = Pt(18)
                r.font.bold = True
            elif level == 2:
                r.font.size = Pt(15)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
            elif level == 3:
                r.font.size = Pt(13)
                r.font.bold = True


def _process_paragraph(container, tag: Tag) -> None:
    p = container.add_paragraph()
    for child in _iter_child_nodes(tag):
        _process_inline(p, child, InlineStyle())


def _process_hr(container) -> None:
    container.add_paragraph("—" * 20)


def _process_block(container, tag: Tag) -> None:
    name = tag.name.lower()

    if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
        _process_heading(container, tag)
        return

    if name == "p":
        _process_paragraph(container, tag)
        return

    if name == "img":
        _add_image(container, tag)
        return

    if name == "hr":
        _process_hr(container)
        return

    if name in ("ul", "ol"):
        _process_list(container, tag, ordered=(name == "ol"))
        return

    if name == "li":
        # fallback: trata como um item bullet
        p = container.add_paragraph(style="List Bullet")
        for child in _iter_child_nodes(tag):
            _process_inline(p, child, InlineStyle())
        return

    if name == "table":
        if _is_real_table(tag):
            _process_table(container, tag)
        else:
            _process_layout_table_as_docx_table(container, tag)
        return

    # default: container
    _process_container(container, tag)


def _process_container(target, container: Tag) -> None:
    for child in _iter_child_nodes(container):
        if isinstance(child, NavigableString):
            text = _normalize_ws(str(child))
            if text:
                target.add_paragraph(text)
            continue

        if isinstance(child, Tag):
            name = child.name.lower()
            if name in BLOCK_TAGS:
                _process_block(target, child)
            else:
                # tags inline soltas dentro de container: vira parágrafo
                p = target.add_paragraph()
                _process_inline(p, child, InlineStyle())


def _html_to_docx_sync(html_content: str) -> BytesIO:
    """
    Converte HTML para DOCX de forma síncrona.

    Args:
        html_content: Conteúdo HTML a ser convertido

    Returns:
        BytesIO com o DOCX gerado.
    """
    global _CSS_CLASS_STYLES
    
    try:
        soup = BeautifulSoup(html_content, "html.parser")

        # captura CSS dos <style> antes de remover
        css_text = "\n".join((t.get_text("\n") or "") for t in soup.find_all("style"))
        _CSS_CLASS_STYLES = _parse_css_class_rules(css_text)

        # remove ruído
        for t in soup.find_all(["script", "style"]):
            t.decompose()

        doc = Document()
        _set_default_document_style(doc)
        
        # respeita @page margin: 1cm (aproximação)
        try:
            section = doc.sections[0]
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)
        except Exception:
            pass

        root = soup.body or soup
        _process_container(doc, root)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    except Exception as e:
        raise HTMLConversionError(f"Erro na conversão HTML para DOCX: {str(e)}")


async def convert_html_to_docx(html_content: str) -> BytesIO:
    """
    Converte HTML para DOCX de forma assíncrona.

    Args:
        html_content: Conteúdo HTML a ser convertido

    Returns:
        BytesIO com o arquivo DOCX
    """
    return await run_in_threadpool(_html_to_docx_sync, html_content)
